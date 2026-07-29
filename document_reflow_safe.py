#!/usr/bin/env python3
"""Safe PDF-to-document conversion and resilient DOCX/PDF regeneration.

The original reflow engine favours semantic text. This layer adds a fidelity
boundary for equations and a validation boundary for models returned by the
browser or an AI provider. Complex source mathematics is preserved as a page
snapshot when extracted Unicode cannot safely represent the original glyphs.
"""
from __future__ import annotations

import base64
import binascii
import copy
import html
import io
import math
import re
import textwrap
import unicodedata
import uuid
from typing import Any

import pymupdf
from docx import Document
from docx.shared import Inches, Pt

import document_reflow as legacy

MODEL_VERSION = 2
MAX_BLOCKS = 5000
MAX_TEXT_CHARS = 250_000
MAX_CELL_CHARS = 25_000
MAX_IMAGE_BYTES = 20 * 1024 * 1024
ALLOWED_TYPES = {
    "paragraph", "heading", "quote", "list_item", "equation",
    "table", "image", "page_break",
}
_IMAGE_RE = re.compile(r"^data:image/(png|jpe?g|webp);base64,(.+)$", re.IGNORECASE | re.DOTALL)
_CONTROL_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")
_MATH_FRAGMENT_RE = re.compile(r"^[\s\dA-Za-z=+\-−×÷*/^_{}\[\]().,;:|<>≤≥≠≈√∑∏∫∞±∓]+$")
_SYMBOL_MAP = str.maketrans({
    "−": "-", "–": "-", "—": "-", "×": "x", "÷": "/",
    "≤": "<=", "≥": ">=", "≠": "!=", "≈": "~=", "∞": "infinity",
    "√": "sqrt", "∑": "sum", "∏": "product", "∫": "integral",
    "→": "->", "←": "<-", "↔": "<->", "⇒": "=>", "⇔": "<=>",
    "λ": "lambda", "μ": "mu", "σ": "sigma", "α": "alpha", "β": "beta",
    "γ": "gamma", "δ": "delta", "θ": "theta", "φ": "phi", "π": "pi",
    "“": '"', "”": '"', "‘": "'", "’": "'", "…": "...",
})


def _id(prefix: str = "block") -> str:
    return f"{prefix}-{uuid.uuid4().hex[:12]}"


def _finite(value: Any, fallback: float, minimum: float, maximum: float) -> float:
    try:
        number = float(value)
    except (TypeError, ValueError):
        return fallback
    if not math.isfinite(number):
        return fallback
    return max(minimum, min(maximum, number))


def _string(value: Any, limit: int = MAX_TEXT_CHARS) -> str:
    text = _CONTROL_RE.sub("", str(value or ""))
    return text[:limit]


def _plain(block: dict[str, Any]) -> str:
    if block.get("type") == "table":
        return "\n".join(" | ".join(map(str, row)) for row in block.get("rows") or [])
    if block.get("type") == "image":
        return _string(block.get("equationText") or block.get("alt") or "")
    return _string(legacy._strip_html(block.get("html") or block.get("text") or ""))


def _unsafe_glyphs(text: str) -> bool:
    if not text:
        return False
    suspicious = 0
    for character in text:
        codepoint = ord(character)
        category = unicodedata.category(character)
        if character == "\ufffd" or 0xE000 <= codepoint <= 0xF8FF or category in {"Cs", "Co"}:
            suspicious += 1
    return suspicious > 0 or suspicious / max(1, len(text)) >= 0.05


def _short_math_fragment(block: dict[str, Any]) -> bool:
    text = _plain(block).strip()
    if not text or len(text) > 28:
        return False
    if re.search(r"[A-Za-z]{4,}", text):
        return False
    return bool(_MATH_FRAGMENT_RE.fullmatch(text)) and (
        any(character.isdigit() for character in text)
        or any(character in "=+−-×÷*/^_{}[]()|≤≥≠≈√∑∏∫∞" for character in text)
    )


def _rect(block: dict[str, Any]) -> list[float] | None:
    value = block.get("rect")
    if not isinstance(value, (list, tuple)) or len(value) != 4:
        return None
    try:
        rect = [float(component) for component in value]
    except (TypeError, ValueError):
        return None
    if not all(math.isfinite(component) for component in rect) or rect[2] <= rect[0] or rect[3] <= rect[1]:
        return None
    return rect


def _source_page(block: dict[str, Any]) -> int | None:
    source = block.get("source")
    if not isinstance(source, dict):
        return None
    try:
        return int(source.get("page"))
    except (TypeError, ValueError):
        return None


def _math_candidate(block: dict[str, Any]) -> bool:
    return (
        block.get("type") == "equation"
        or _unsafe_glyphs(_plain(block))
        or _short_math_fragment(block)
    ) and _rect(block) is not None and _source_page(block) is not None


def _blocks_near(first: dict[str, Any], second: dict[str, Any]) -> bool:
    first_rect = _rect(first)
    second_rect = _rect(second)
    if not first_rect or not second_rect or _source_page(first) != _source_page(second):
        return False
    vertical_gap = max(0.0, second_rect[1] - first_rect[3])
    horizontal_gap = max(0.0, max(first_rect[0], second_rect[0]) - min(first_rect[2], second_rect[2]))
    height = max(first_rect[3] - first_rect[1], second_rect[3] - second_rect[1], 8.0)
    return vertical_gap <= max(34.0, height * 2.1) and horizontal_gap <= max(110.0, height * 8.0)


def _snapshot_group(document: Any, group: list[dict[str, Any]]) -> dict[str, Any] | None:
    page_number = _source_page(group[0])
    rects = [_rect(block) for block in group]
    if page_number is None or any(rect is None for rect in rects):
        return None
    page = document.load_page(page_number)
    valid_rects = [rect for rect in rects if rect is not None]
    clip = pymupdf.Rect(
        min(rect[0] for rect in valid_rects) - 5,
        min(rect[1] for rect in valid_rects) - 5,
        max(rect[2] for rect in valid_rects) + 5,
        max(rect[3] for rect in valid_rects) + 5,
    ) & page.rect
    if clip.is_empty or clip.width < 2 or clip.height < 2:
        return None
    pixmap = page.get_pixmap(matrix=pymupdf.Matrix(2.5, 2.5), clip=clip, alpha=False)
    png = pixmap.tobytes("png")
    source_text = "\n".join(filter(None, (_plain(block).strip() for block in group)))
    return {
        "id": _id("equation-image"),
        "type": "image",
        "rect": [clip.x0, clip.y0, clip.x1, clip.y1],
        "dataUrl": f"data:image/png;base64,{base64.b64encode(png).decode('ascii')}",
        "mime": "image/png",
        "width": max(80.0, min(float(page.rect.width) * 0.84, float(clip.width))),
        "height": max(24.0, float(clip.height)),
        "alt": f"Preserved equation from page {page_number + 1}",
        "equationText": source_text,
        "source": {
            "kind": "equation_snapshot",
            "page": page_number,
            "rect": [clip.x0, clip.y0, clip.x1, clip.y1],
            "members": [str(block.get("id") or "") for block in group],
        },
    }


def preserve_complex_math(pdf_bytes: bytes, model: dict[str, Any]) -> tuple[dict[str, Any], list[str]]:
    """Replace unsafe source equation runs with faithful page snapshots."""
    result = copy.deepcopy(model)
    blocks = list(result.get("blocks") or [])
    warnings: list[str] = []
    document = pymupdf.open(stream=pdf_bytes, filetype="pdf")
    try:
        output: list[dict[str, Any]] = []
        index = 0
        while index < len(blocks):
            block = blocks[index]
            if not _math_candidate(block):
                output.append(block)
                index += 1
                continue
            group = [block]
            cursor = index + 1
            while cursor < len(blocks) and _math_candidate(blocks[cursor]) and _blocks_near(group[-1], blocks[cursor]):
                group.append(blocks[cursor])
                cursor += 1
            must_snapshot = any(item.get("type") == "equation" or _unsafe_glyphs(_plain(item)) for item in group)
            numeric_rows = sum(any(character.isdigit() for character in _plain(item)) for item in group)
            if must_snapshot or len(group) >= 3 and numeric_rows >= 2:
                snapshot = _snapshot_group(document, group)
                if snapshot:
                    output.append(snapshot)
                    warnings.append(
                        f"Preserved complex mathematics from page {_source_page(group[0]) + 1} as an image because its PDF glyph encoding was not reliable."
                    )
                    index = cursor
                    continue
            output.extend(group)
            index = cursor
    finally:
        document.close()
    result["blocks"] = output
    result["warnings"] = list(result.get("warnings") or []) + warnings
    return result, warnings


def _normalise_style(value: Any) -> dict[str, Any]:
    style = value if isinstance(value, dict) else {}
    color = str(style.get("color") or "#111318")
    if not re.fullmatch(r"#[0-9a-fA-F]{6}", color):
        color = "#111318"
    align = str(style.get("align") or "left").lower()
    if align not in {"left", "center", "right", "justify"}:
        align = "left"
    return {
        "fontFamily": _string(style.get("fontFamily") or "Arial", 120),
        "fontSize": _finite(style.get("fontSize"), 11.0, 5.0, 96.0),
        "bold": bool(style.get("bold", False)),
        "italic": bool(style.get("italic", False)),
        "underline": bool(style.get("underline", False)),
        "color": color,
        "align": align,
        "lineHeight": _finite(style.get("lineHeight"), 1.2, 0.8, 3.0),
    }


def _normalise_image(block: dict[str, Any]) -> dict[str, Any] | None:
    data_url = _string(block.get("dataUrl"), MAX_IMAGE_BYTES * 2)
    match = _IMAGE_RE.match(data_url)
    if not match:
        return None
    try:
        raw = base64.b64decode(match.group(2), validate=False)
    except (binascii.Error, ValueError):
        return None
    if not raw or len(raw) > MAX_IMAGE_BYTES:
        return None
    mime = match.group(1).lower()
    if mime == "jpg":
        mime = "jpeg"
    if mime == "webp":
        try:
            image_document = pymupdf.open(stream=raw)
            try:
                pixmap = image_document[0].get_pixmap(alpha=False)
                raw = pixmap.tobytes("png")
                mime = "png"
            finally:
                image_document.close()
        except Exception:
            return None
    return {
        "id": _string(block.get("id") or _id("image"), 160),
        "type": "image",
        "dataUrl": f"data:image/{mime};base64,{base64.b64encode(raw).decode('ascii')}",
        "mime": f"image/{mime}",
        "width": _finite(block.get("width"), 320.0, 24.0, 1600.0),
        "height": _finite(block.get("height"), 200.0, 16.0, 1600.0),
        "alt": _string(block.get("alt") or "Document image", 1000),
        "equationText": _string(block.get("equationText"), 20_000),
        "source": copy.deepcopy(block.get("source")) if isinstance(block.get("source"), dict) else {},
    }


def sanitise_model(model: Any) -> tuple[dict[str, Any], list[str]]:
    warnings: list[str] = []
    source = model if isinstance(model, dict) else {}
    page_value = source.get("page") if isinstance(source.get("page"), dict) else {}
    width = _finite(page_value.get("width"), 612.0, 216.0, 2000.0)
    height = _finite(page_value.get("height"), 792.0, 216.0, 3000.0)
    margin = _finite(page_value.get("margin"), 54.0, 12.0, min(width, height) * 0.35)
    clean: dict[str, Any] = {
        "version": MODEL_VERSION,
        "title": _string(source.get("title") or "Lumina document", 1000),
        "page": {"width": width, "height": height, "margin": margin},
        "blocks": [],
        "warnings": list(source.get("warnings") or [])[:200],
        "source": copy.deepcopy(source.get("source")) if isinstance(source.get("source"), dict) else {},
    }
    raw_blocks = source.get("blocks") if isinstance(source.get("blocks"), list) else []
    if len(raw_blocks) > MAX_BLOCKS:
        warnings.append(f"The document contained more than {MAX_BLOCKS} blocks; extra blocks were ignored.")
    for raw in raw_blocks[:MAX_BLOCKS]:
        if not isinstance(raw, dict):
            warnings.append("Ignored a malformed document block.")
            continue
        block_type = str(raw.get("type") or "paragraph")
        if block_type not in ALLOWED_TYPES:
            block_type = "paragraph"
        if block_type == "page_break":
            clean["blocks"].append({"id": _string(raw.get("id") or _id("break"), 160), "type": "page_break"})
            continue
        if block_type == "image":
            image_block = _normalise_image(raw)
            if image_block:
                clean["blocks"].append(image_block)
            else:
                warnings.append(f"Replaced an unreadable image block with a text placeholder: {_string(raw.get('alt') or 'image', 120)}.")
                placeholder = _string(raw.get("equationText") or raw.get("alt") or "[Image unavailable]")
                clean["blocks"].append({
                    "id": _string(raw.get("id") or _id("image-placeholder"), 160),
                    "type": "paragraph",
                    "text": placeholder,
                    "html": html.escape(placeholder),
                    "style": _normalise_style(raw.get("style")),
                })
            continue
        if block_type == "table":
            rows: list[list[str]] = []
            for raw_row in (raw.get("rows") if isinstance(raw.get("rows"), list) else [])[:200]:
                if not isinstance(raw_row, list):
                    raw_row = [raw_row]
                rows.append([_string(cell, MAX_CELL_CHARS) for cell in raw_row[:50]])
            if not rows:
                rows = [[""]]
            clean["blocks"].append({
                "id": _string(raw.get("id") or _id("table"), 160),
                "type": "table",
                "rows": rows,
                "style": _normalise_style(raw.get("style")),
            })
            continue
        text = _string(raw.get("text") or legacy._strip_html(raw.get("html") or ""))
        markup = _string(raw.get("html") or html.escape(text), MAX_TEXT_CHARS * 2)
        block: dict[str, Any] = {
            "id": _string(raw.get("id") or _id(block_type), 160),
            "type": block_type,
            "text": text,
            "html": markup,
            "style": _normalise_style(raw.get("style")),
        }
        if block_type == "heading":
            block["level"] = int(_finite(raw.get("level"), 2, 1, 6))
        if block_type == "list_item":
            block["listType"] = "number" if raw.get("listType") == "number" else "bullet"
        if block_type == "equation":
            block["latex"] = _string(raw.get("latex") or text, MAX_TEXT_CHARS)
            block["style"]["fontFamily"] = "Cambria Math"
            block["style"]["align"] = "center"
        clean["blocks"].append(block)
    clean["warnings"].extend(warnings)
    return clean, warnings


def pdf_to_document_model(pdf_bytes: bytes, title: str = "Document") -> dict[str, Any]:
    imported = legacy.pdf_to_document_model(pdf_bytes, title=title)
    preserved, _warnings = preserve_complex_math(pdf_bytes, imported)
    clean, sanitise_warnings = sanitise_model(preserved)
    clean["warnings"].extend(sanitise_warnings)
    return clean


def _pdf_text(value: Any) -> str:
    text = unicodedata.normalize("NFKC", _string(value)).translate(_SYMBOL_MAP)
    filtered: list[str] = []
    for character in text:
        codepoint = ord(character)
        category = unicodedata.category(character)
        if character in "\n\t":
            filtered.append(character)
        elif character == "\ufffd" or 0xE000 <= codepoint <= 0xF8FF or category in {"Cs", "Co", "Cc"}:
            filtered.append("?")
        else:
            filtered.append(character)
    return "".join(filtered).encode("cp1252", errors="replace").decode("cp1252")


def _fallback_model(model: dict[str, Any]) -> dict[str, Any]:
    result = copy.deepcopy(model)
    for block in result.get("blocks") or []:
        if block.get("type") == "table":
            block["rows"] = [[_pdf_text(cell) for cell in row] for row in block.get("rows") or []]
        elif block.get("type") not in {"image", "page_break"}:
            text = _pdf_text(legacy._strip_html(block.get("html") or block.get("text") or ""))
            block["text"] = text
            block["html"] = html.escape(text).replace("\n", "<br>")
    return result


def _emergency_docx(model: dict[str, Any]) -> bytes:
    document = Document()
    document.core_properties.title = str(model.get("title") or "Lumina document")
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    for block in model.get("blocks") or []:
        block_type = block.get("type")
        if block_type == "page_break":
            document.add_page_break()
        elif block_type == "image":
            data_url = str(block.get("dataUrl") or "")
            match = _IMAGE_RE.match(data_url)
            if match:
                try:
                    raw = base64.b64decode(match.group(2), validate=False)
                    document.add_picture(io.BytesIO(raw), width=Inches(min(6.5, float(block.get("width") or 320) / 72.0)))
                except Exception:
                    document.add_paragraph(f"[{block.get('alt') or 'Image'}]")
        elif block_type == "table":
            rows = block.get("rows") or [[""]]
            columns = max(len(row) for row in rows)
            table = document.add_table(rows=len(rows), cols=columns)
            table.style = "Table Grid"
            for row_index, row in enumerate(rows):
                for column_index in range(columns):
                    table.cell(row_index, column_index).text = str(row[column_index] if column_index < len(row) else "")
        else:
            text = legacy._strip_html(block.get("html") or block.get("text") or "")
            paragraph = document.add_heading(text, level=max(1, min(6, int(block.get("level") or 2)))) if block_type == "heading" else document.add_paragraph(text)
            if block_type == "equation":
                paragraph.alignment = 1
    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def _emergency_pdf(model: dict[str, Any]) -> bytes:
    page_config = model.get("page") or {}
    width = _finite(page_config.get("width"), 612, 216, 2000)
    height = _finite(page_config.get("height"), 792, 216, 3000)
    margin = _finite(page_config.get("margin"), 54, 12, min(width, height) * 0.35)
    document = pymupdf.open()
    page = document.new_page(width=width, height=height)
    y = margin

    def next_page() -> None:
        nonlocal page, y
        page = document.new_page(width=width, height=height)
        y = margin

    for block in model.get("blocks") or []:
        if block.get("type") == "page_break":
            next_page()
            continue
        if block.get("type") == "image":
            label = _pdf_text(block.get("alt") or "[Image]")
            lines = [label]
        elif block.get("type") == "table":
            lines = [_pdf_text(" | ".join(map(str, row))) for row in block.get("rows") or []]
        else:
            raw = legacy._strip_html(block.get("html") or block.get("text") or "")
            lines = []
            for paragraph in _pdf_text(raw).splitlines() or [""]:
                lines.extend(textwrap.wrap(paragraph, width=92) or [""])
        font_size = 14 if block.get("type") == "heading" else 10
        for line in lines:
            if y + font_size * 1.5 > height - margin:
                next_page()
            page.insert_text((margin, y), line[:500], fontsize=font_size, fontname="helv")
            y += font_size * 1.45
        y += 5
    output = document.tobytes(garbage=4, deflate=True, clean=True)
    document.close()
    return output


def render_document_model(model: Any, prefer_office: bool = True) -> dict[str, Any]:
    clean, warnings = sanitise_model(model)
    recovered = False
    try:
        docx_bytes = legacy.model_to_docx_bytes(clean)
    except Exception as error:
        recovered = True
        warnings.append(f"DOCX generation recovered from a malformed block: {error}")
        docx_bytes = _emergency_docx(clean)

    pdf_bytes: bytes | None = None
    converter = "pymupdf-reflow-safe"
    if prefer_office:
        try:
            pdf_bytes, office_converter = legacy.convert_docx_bytes_to_pdf(docx_bytes)
            if pdf_bytes:
                converter = office_converter or "office"
        except Exception as error:
            recovered = True
            warnings.append(f"Office conversion was unavailable; Lumina used its local PDF renderer: {error}")
    if not pdf_bytes:
        try:
            pdf_bytes = legacy.model_to_fallback_pdf_bytes(_fallback_model(clean))
        except Exception as error:
            recovered = True
            warnings.append(f"The rich PDF renderer recovered using an emergency readable layout: {error}")
            converter = "pymupdf-emergency"
            pdf_bytes = _emergency_pdf(clean)

    try:
        markdown = legacy.model_to_markdown(clean)
    except Exception as error:
        recovered = True
        warnings.append(f"Markdown export omitted unsupported formatting: {error}")
        markdown = "\n\n".join(_plain(block) for block in clean.get("blocks") or [])

    return {
        "model": clean,
        "docx": docx_bytes,
        "pdf": pdf_bytes,
        "markdown": markdown,
        "converter": converter,
        "warnings": warnings,
        "recovered": recovered,
    }


__all__ = [
    "MODEL_VERSION",
    "pdf_to_document_model",
    "preserve_complex_math",
    "sanitise_model",
    "render_document_model",
]
