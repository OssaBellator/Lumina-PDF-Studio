#!/usr/bin/env python3
"""PDF <-> editable document model, DOCX, Markdown, and reflowed PDF helpers."""
from __future__ import annotations

import base64
import html
import io
import os
import re
import shutil
import subprocess
import tempfile
import uuid
from html.parser import HTMLParser
from pathlib import Path
from typing import Any

try:
    import pymupdf
except ImportError:  # pragma: no cover
    import fitz as pymupdf

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

MODEL_VERSION = 1
DEFAULT_PAGE = {"width": 612.0, "height": 792.0, "margin": 54.0}
MATH_CHARS = set("=+-−×÷∑∏∫√∞≈≠≤≥∂∇∈∉⊂⊃⊆⊇∪∩→←↔⇒⇔λμσαβγδθφψωΓΔΘΛΞΠΣΦΨΩ^_{}[]()|±∓")
LIST_RE = re.compile(r"^\s*(?:[-*•]|\d+[.)]|[A-Za-z][.)])\s+")


def _id(prefix: str = "block") -> str:
    return f"{prefix}-{uuid.uuid4().hex[:12]}"


def _color_hex(value: int | None) -> str:
    red, green, blue = pymupdf.sRGB_to_pdf(int(value or 0))
    return "#" + "".join(f"{max(0, min(255, round(component * 255))):02x}" for component in (red, green, blue))


def _escape_text(value: Any) -> str:
    return html.escape(str(value or "")).replace("\n", "<br>")


def _strip_html(value: Any) -> str:
    parser = _PlainTextParser()
    parser.feed(str(value or ""))
    return parser.text.strip("\n")


class _PlainTextParser(HTMLParser):
    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: list[str] = []

    @property
    def text(self) -> str:
        return "".join(self.parts)

    def handle_data(self, data: str) -> None:
        self.parts.append(data)

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        if tag in {"br", "p", "div", "li"} and self.parts and not self.parts[-1].endswith("\n"):
            self.parts.append("\n")

    def handle_endtag(self, tag: str) -> None:
        if tag in {"p", "div", "li"} and self.parts and not self.parts[-1].endswith("\n"):
            self.parts.append("\n")


class _RunParser(HTMLParser):
    """Parse a deliberately small rich-text subset into DOCX-friendly runs."""

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.stack: list[dict[str, Any]] = [{}]
        self.runs: list[tuple[str, dict[str, Any]]] = []

    def handle_starttag(self, tag: str, attrs: list[tuple[str, str | None]]) -> None:
        current = dict(self.stack[-1])
        if tag in {"b", "strong"}:
            current["bold"] = True
        elif tag in {"i", "em"}:
            current["italic"] = True
        elif tag == "u":
            current["underline"] = True
        elif tag == "code":
            current["font"] = "Courier New"
        elif tag == "span":
            style = dict(attrs).get("style") or ""
            color = re.search(r"color\s*:\s*(#[0-9a-fA-F]{6})", style)
            size = re.search(r"font-size\s*:\s*([0-9.]+)px", style)
            family = re.search(r"font-family\s*:\s*([^;]+)", style)
            if color:
                current["color"] = color.group(1)
            if size:
                current["size"] = float(size.group(1)) * 0.75
            if family:
                current["font"] = family.group(1).strip(" '\"")
        elif tag == "br":
            self.runs.append(("\n", dict(current)))
        self.stack.append(current)

    def handle_endtag(self, tag: str) -> None:
        if len(self.stack) > 1:
            self.stack.pop()

    def handle_data(self, data: str) -> None:
        if data:
            self.runs.append((data, dict(self.stack[-1])))


def _rect(item: dict[str, Any]) -> list[float]:
    values = item.get("rect") or [0, 0, 0, 0]
    return [float(values[0]), float(values[1]), float(values[2]), float(values[3])]


def _inside(rect: list[float], container: list[float], tolerance: float = 1.0) -> bool:
    cx = (rect[0] + rect[2]) / 2
    cy = (rect[1] + rect[3]) / 2
    return container[0] - tolerance <= cx <= container[2] + tolerance and container[1] - tolerance <= cy <= container[3] + tolerance


def _looks_math(text: str, font_name: str = "") -> bool:
    compact = "".join(text.split())
    if not compact:
        return False
    marker = any(token in font_name.lower() for token in ("math", "symbol", "cmmi", "cmsy", "cmex", "stix"))
    symbols = sum(character in MATH_CHARS or 0x2200 <= ord(character) <= 0x22FF for character in compact)
    numbers = sum(character.isdigit() for character in compact)
    return marker or symbols >= 2 or (symbols >= 1 and numbers >= 1 and len(compact) <= 72)


def _line_style(line: dict[str, Any]) -> dict[str, Any]:
    spans = list(line.get("spans") or [])
    dominant = max(spans, key=lambda span: len(str(span.get("text") or "")), default={})
    flags = int(dominant.get("flags") or 0)
    return {
        "fontFamily": str(dominant.get("font") or "Arial"),
        "fontSize": float(dominant.get("size") or 11),
        "bold": bool(flags & getattr(pymupdf, "TEXT_FONT_BOLD", 16)),
        "italic": bool(flags & getattr(pymupdf, "TEXT_FONT_ITALIC", 2)),
        "underline": False,
        "color": _color_hex(dominant.get("color", 0)),
        "align": "left",
        "lineHeight": 1.15,
    }


def _extract_tables(page: Any) -> list[dict[str, Any]]:
    tables: list[dict[str, Any]] = []
    finder = getattr(page, "find_tables", None)
    if finder is None:
        return tables
    try:
        found = finder()
        for index, table in enumerate(getattr(found, "tables", []) or []):
            bbox = list(table.bbox)
            rows = [[str(cell or "") for cell in row] for row in (table.extract() or [])]
            if rows and any(any(cell.strip() for cell in row) for row in rows):
                tables.append({
                    "id": _id("table"),
                    "type": "table",
                    "rect": bbox,
                    "rows": rows,
                    "style": {"fontFamily": "Arial", "fontSize": 10, "color": "#111318"},
                    "source": {"kind": "table", "index": index},
                })
    except Exception:
        return []
    return tables


def _extract_page_items(page: Any, page_index: int) -> list[dict[str, Any]]:
    tables = _extract_tables(page)
    table_rects = [_rect(table) for table in tables]
    data = page.get_text("dict", flags=pymupdf.TEXTFLAGS_DICT | pymupdf.TEXT_PRESERVE_IMAGES | pymupdf.TEXT_PRESERVE_WHITESPACE)
    items: list[dict[str, Any]] = list(tables)
    lines: list[dict[str, Any]] = []
    for block_index, block in enumerate(data.get("blocks", [])):
        block_type = int(block.get("type", -1))
        if block_type == 0:
            for line_index, line in enumerate(block.get("lines", [])):
                text = "".join(str(span.get("text") or "") for span in line.get("spans", [])).strip()
                if not text:
                    continue
                rect = list(line.get("bbox") or block.get("bbox") or [0, 0, 0, 0])
                if any(_inside(rect, table_rect) for table_rect in table_rects):
                    continue
                style = _line_style(line)
                lines.append({
                    "id": f"p{page_index}-line-{block_index}-{line_index}",
                    "text": text,
                    "rect": [float(value) for value in rect],
                    "style": style,
                    "math": _looks_math(text, style["fontFamily"]),
                })
        elif block_type == 1:
            raw = bytes(block.get("image") or b"")
            if not raw:
                continue
            extension = str(block.get("ext") or "png").lower()
            mime = "image/jpeg" if extension in {"jpg", "jpeg"} else f"image/{extension}"
            rect = [float(value) for value in (block.get("bbox") or [0, 0, 0, 0])]
            items.append({
                "id": _id("image"),
                "type": "image",
                "rect": rect,
                "dataUrl": f"data:{mime};base64,{base64.b64encode(raw).decode('ascii')}",
                "mime": mime,
                "width": max(24.0, rect[2] - rect[0]),
                "height": max(24.0, rect[3] - rect[1]),
                "alt": f"Image from page {page_index + 1}",
                "source": {"kind": "image", "page": page_index},
            })

    lines.sort(key=lambda item: (item["rect"][1], item["rect"][0]))
    body_sizes = sorted(item["style"]["fontSize"] for item in lines if not item["math"])
    body_size = body_sizes[(len(body_sizes) - 1) // 2] if body_sizes else 11.0
    paragraph_lines: list[dict[str, Any]] = []
    math_buffer: list[dict[str, Any]] = []

    def flush_math() -> None:
        nonlocal math_buffer
        if not math_buffer:
            return
        rects = [entry["rect"] for entry in math_buffer]
        items.append({
            "id": _id("equation"),
            "type": "equation",
            "rect": [min(r[0] for r in rects), min(r[1] for r in rects), max(r[2] for r in rects), max(r[3] for r in rects)],
            "html": _escape_text("\n".join(entry["text"] for entry in math_buffer)),
            "text": "\n".join(entry["text"] for entry in math_buffer),
            "latex": "",
            "style": {**math_buffer[0]["style"], "fontFamily": "Cambria Math", "align": "center"},
            "source": {"kind": "math", "page": page_index, "members": [entry["id"] for entry in math_buffer]},
        })
        math_buffer = []

    def flush_paragraph() -> None:
        nonlocal paragraph_lines
        if not paragraph_lines:
            return
        rects = [entry["rect"] for entry in paragraph_lines]
        first = paragraph_lines[0]
        text = " ".join(entry["text"].strip() for entry in paragraph_lines).strip()
        style = dict(first["style"])
        short = len(text) <= 120
        is_heading = style["fontSize"] >= body_size * 1.28 or (style["bold"] and short and style["fontSize"] >= body_size * 1.05)
        list_match = LIST_RE.match(text)
        block_type = "heading" if is_heading else "list_item" if list_match else "paragraph"
        level = 1 if style["fontSize"] >= body_size * 1.75 else 2 if style["fontSize"] >= body_size * 1.42 else 3
        items.append({
            "id": _id("text"),
            "type": block_type,
            "level": level if is_heading else None,
            "listType": "number" if list_match and re.match(r"^\s*\d", text) else "bullet" if list_match else None,
            "rect": [min(r[0] for r in rects), min(r[1] for r in rects), max(r[2] for r in rects), max(r[3] for r in rects)],
            "html": _escape_text(text),
            "text": text,
            "style": style,
            "source": {"kind": "text", "page": page_index, "members": [entry["id"] for entry in paragraph_lines]},
        })
        paragraph_lines = []

    for line in lines:
        if line["math"]:
            flush_paragraph()
            if math_buffer:
                last = math_buffer[-1]
                gap = line["rect"][1] - last["rect"][3]
                horizontal_near = line["rect"][0] <= max(last["rect"][2] + 36, last["rect"][0] + 180)
                if gap > max(22.0, line["style"]["fontSize"] * 1.8) or not horizontal_near:
                    flush_math()
            math_buffer.append(line)
            continue
        flush_math()
        if paragraph_lines:
            last = paragraph_lines[-1]
            gap = line["rect"][1] - last["rect"][3]
            same_column = abs(line["rect"][0] - last["rect"][0]) <= max(18.0, body_size * 2)
            similar_size = abs(line["style"]["fontSize"] - last["style"]["fontSize"]) <= 1.5
            if gap > max(10.0, body_size * 0.9) or not same_column or not similar_size:
                flush_paragraph()
        paragraph_lines.append(line)
    flush_paragraph()
    flush_math()
    return sorted(items, key=lambda item: (_rect(item)[1], _rect(item)[0]))


def pdf_to_document_model(pdf_bytes: bytes, title: str = "Document") -> dict[str, Any]:
    document = pymupdf.open(stream=pdf_bytes, filetype="pdf")
    try:
        if not document.page_count:
            return {"version": MODEL_VERSION, "title": title, "page": dict(DEFAULT_PAGE), "blocks": [], "warnings": ["The PDF has no pages."]}
        first = document.load_page(0)
        page_config = {"width": float(first.rect.width), "height": float(first.rect.height), "margin": max(36.0, min(72.0, first.rect.width * 0.08))}
        blocks: list[dict[str, Any]] = []
        warnings: list[str] = []
        for page_index in range(document.page_count):
            page = document.load_page(page_index)
            if page_index:
                blocks.append({"id": _id("break"), "type": "page_break"})
            page_items = _extract_page_items(page, page_index)
            if not page_items:
                warnings.append(f"Page {page_index + 1} did not contain editable text, tables, or embedded images.")
            blocks.extend(page_items)
        return {
            "version": MODEL_VERSION,
            "title": title,
            "page": page_config,
            "blocks": blocks,
            "warnings": warnings,
            "source": {"format": "pdf", "pageCount": document.page_count},
        }
    finally:
        document.close()


def _apply_run_style(run: Any, style: dict[str, Any]) -> None:
    run.bold = bool(style.get("bold", False))
    run.italic = bool(style.get("italic", False))
    run.underline = bool(style.get("underline", False))
    if style.get("font"):
        run.font.name = str(style["font"])
    if style.get("size"):
        run.font.size = Pt(float(style["size"]))
    color = str(style.get("color") or "")
    if re.fullmatch(r"#[0-9a-fA-F]{6}", color):
        run.font.color.rgb = RGBColor.from_string(color[1:].upper())


def _paragraph_alignment(value: Any) -> WD_ALIGN_PARAGRAPH:
    return {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(str(value or "left").lower(), WD_ALIGN_PARAGRAPH.LEFT)


def _add_html_to_paragraph(paragraph: Any, markup: str, block_style: dict[str, Any]) -> None:
    parser = _RunParser()
    parser.feed(markup or "")
    runs = parser.runs or [("", {})]
    for text, inline in runs:
        pieces = text.split("\n")
        for index, piece in enumerate(pieces):
            if index:
                paragraph.add_run().add_break()
            if piece:
                run = paragraph.add_run(piece)
                merged = {
                    "bold": block_style.get("bold", False),
                    "italic": block_style.get("italic", False),
                    "underline": block_style.get("underline", False),
                    "font": block_style.get("fontFamily") or "Arial",
                    "size": block_style.get("fontSize") or 11,
                    "color": block_style.get("color") or "#111318",
                    **inline,
                }
                _apply_run_style(run, merged)


def _set_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def model_to_docx_bytes(model: dict[str, Any]) -> bytes:
    document = Document()
    section = document.sections[0]
    page = {**DEFAULT_PAGE, **(model.get("page") or {})}
    section.page_width = Inches(float(page["width"]) / 72.0)
    section.page_height = Inches(float(page["height"]) / 72.0)
    margin = Inches(float(page["margin"]) / 72.0)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = margin
    document.core_properties.title = str(model.get("title") or "Lumina document")
    document.core_properties.subject = "Editable document generated from a PDF by Lumina PDF Studio"
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    for block in model.get("blocks") or []:
        block_type = str(block.get("type") or "paragraph")
        style = dict(block.get("style") or {})
        if block_type == "page_break":
            document.add_page_break()
            continue
        if block_type == "table":
            rows = block.get("rows") or []
            column_count = max((len(row) for row in rows), default=1)
            table = document.add_table(rows=max(1, len(rows)), cols=max(1, column_count))
            table.style = "Table Grid"
            for row_index, row in enumerate(rows):
                for column_index in range(column_count):
                    value = str(row[column_index] if column_index < len(row) else "")
                    cell = table.cell(row_index, column_index)
                    cell.text = value
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = str(style.get("fontFamily") or "Arial")
                            run.font.size = Pt(float(style.get("fontSize") or 10))
                            if row_index == 0:
                                run.bold = True
                    if row_index == 0:
                        _set_cell_shading(cell, "EDE9FE")
            continue
        if block_type == "image":
            data_url = str(block.get("dataUrl") or "")
            if data_url.startswith("data:") and "," in data_url:
                try:
                    raw = base64.b64decode(data_url.split(",", 1)[1])
                    width_points = min(float(block.get("width") or 320), float(page["width"]) - float(page["margin"]) * 2)
                    document.add_picture(io.BytesIO(raw), width=Inches(max(0.5, width_points / 72.0)))
                    paragraph = document.paragraphs[-1]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    paragraph = document.add_paragraph(f"[Image unavailable: {block.get('alt') or 'image'}]")
                    paragraph.style = document.styles["Caption"]
            continue

        text_html = str(block.get("html") or _escape_text(block.get("text") or ""))
        if block_type == "heading":
            paragraph = document.add_heading(level=max(1, min(6, int(block.get("level") or 2))))
        elif block_type == "list_item":
            paragraph = document.add_paragraph(style="List Number" if block.get("listType") == "number" else "List Bullet")
        elif block_type == "quote":
            paragraph = document.add_paragraph(style="Intense Quote")
        else:
            paragraph = document.add_paragraph()
        paragraph.alignment = _paragraph_alignment(style.get("align"))
        paragraph.paragraph_format.line_spacing = float(style.get("lineHeight") or 1.15)
        if block_type == "equation":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            style = {**style, "fontFamily": "Cambria Math"}
        _add_html_to_paragraph(paragraph, text_html, style)

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


def model_to_markdown(model: dict[str, Any]) -> str:
    parts: list[str] = []
    for block in model.get("blocks") or []:
        block_type = str(block.get("type") or "paragraph")
        text = _strip_html(block.get("html") or block.get("text") or "")
        if block_type == "page_break":
            parts.append("\n---\n")
        elif block_type == "heading":
            parts.append(f"{'#' * max(1, min(6, int(block.get('level') or 2)))} {text}")
        elif block_type == "list_item":
            marker = "1." if block.get("listType") == "number" else "-"
            parts.append(f"{marker} {text}")
        elif block_type == "equation":
            latex = str(block.get("latex") or "").strip()
            parts.append(f"$$\n{latex or text}\n$$")
        elif block_type == "quote":
            parts.append("\n".join(f"> {line}" for line in text.splitlines()))
        elif block_type == "table":
            rows = block.get("rows") or []
            if rows:
                width = max(len(row) for row in rows)
                normalized = [list(row) + [""] * (width - len(row)) for row in rows]
                parts.append("| " + " | ".join(normalized[0]) + " |")
                parts.append("| " + " | ".join(["---"] * width) + " |")
                parts.extend("| " + " | ".join(row) + " |" for row in normalized[1:])
        elif block_type == "image":
            parts.append(f"![{block.get('alt') or 'Image'}]({block.get('dataUrl') or ''})")
        else:
            parts.append(text)
    return "\n\n".join(part for part in parts if part is not None).strip() + "\n"


def _font_alias(style: dict[str, Any]) -> str:
    family = str(style.get("fontFamily") or "Helvetica").lower()
    bold = bool(style.get("bold"))
    italic = bool(style.get("italic"))
    if "cour" in family or "mono" in family:
        return "cobi" if bold and italic else "cobo" if bold else "coit" if italic else "cour"
    if "times" in family or "serif" in family or "cambria" in family:
        return "tibi" if bold and italic else "tibo" if bold else "tiit" if italic else "tiro"
    return "hebi" if bold and italic else "hebo" if bold else "heit" if italic else "helv"


def _rgb(value: Any) -> tuple[float, float, float]:
    text = str(value or "#111318")
    if re.fullmatch(r"#[0-9a-fA-F]{6}", text):
        return tuple(int(text[index:index + 2], 16) / 255 for index in (1, 3, 5))  # type: ignore[return-value]
    return (0.07, 0.075, 0.095)


def _wrap_text(text: str, width: float, font: str, size: float) -> list[str]:
    result: list[str] = []
    for paragraph in text.splitlines() or [""]:
        words = paragraph.split()
        if not words:
            result.append("")
            continue
        line = words[0]
        for word in words[1:]:
            candidate = f"{line} {word}"
            if pymupdf.get_text_length(candidate, fontname=font, fontsize=size) <= width:
                line = candidate
            else:
                result.append(line)
                line = word
        result.append(line)
    return result


def model_to_fallback_pdf_bytes(model: dict[str, Any]) -> bytes:
    page_config = {**DEFAULT_PAGE, **(model.get("page") or {})}
    width = float(page_config["width"])
    height = float(page_config["height"])
    margin = float(page_config["margin"])
    document = pymupdf.open()
    page = document.new_page(width=width, height=height)
    y = margin

    def new_page() -> Any:
        nonlocal page, y
        page = document.new_page(width=width, height=height)
        y = margin
        return page

    def ensure_space(required: float) -> None:
        if y + required > height - margin:
            new_page()

    for block in model.get("blocks") or []:
        block_type = str(block.get("type") or "paragraph")
        if block_type == "page_break":
            new_page()
            continue
        if block_type == "image":
            data_url = str(block.get("dataUrl") or "")
            if data_url.startswith("data:") and "," in data_url:
                try:
                    raw = base64.b64decode(data_url.split(",", 1)[1])
                    target_width = min(float(block.get("width") or 300), width - margin * 2)
                    aspect = max(0.1, float(block.get("width") or 300) / max(1.0, float(block.get("height") or 200)))
                    target_height = target_width / aspect
                    ensure_space(target_height + 12)
                    rect = pymupdf.Rect(margin, y, margin + target_width, y + target_height)
                    page.insert_image(rect, stream=raw, keep_proportion=True)
                    y += target_height + 12
                except Exception:
                    pass
            continue
        if block_type == "table":
            rows = block.get("rows") or []
            if not rows:
                continue
            columns = max(len(row) for row in rows)
            cell_width = (width - margin * 2) / max(1, columns)
            style = block.get("style") or {}
            font = _font_alias(style)
            size = float(style.get("fontSize") or 9)
            for row_index, row in enumerate(rows):
                wrapped = [_wrap_text(str(row[col] if col < len(row) else ""), cell_width - 8, font, size) for col in range(columns)]
                row_height = max(22.0, max(len(lines) for lines in wrapped) * size * 1.3 + 8)
                ensure_space(row_height)
                for col, lines in enumerate(wrapped):
                    rect = pymupdf.Rect(margin + col * cell_width, y, margin + (col + 1) * cell_width, y + row_height)
                    page.draw_rect(rect, color=(0.65, 0.65, 0.7), fill=(0.94, 0.93, 0.99) if row_index == 0 else None, width=0.5)
                    page.insert_textbox(rect + (4, 3, -4, -3), "\n".join(lines), fontname=font, fontsize=size, color=_rgb(style.get("color")), align=0)
                y += row_height
            y += 12
            continue

        style = dict(block.get("style") or {})
        text = _strip_html(block.get("html") or block.get("text") or "")
        if block_type == "heading":
            level = max(1, min(6, int(block.get("level") or 2)))
            style["fontSize"] = max(float(style.get("fontSize") or 11), {1: 24, 2: 20, 3: 16}.get(level, 13))
            style["bold"] = True
        elif block_type == "equation":
            style["fontFamily"] = "Times"
            style["align"] = "center"
        font = _font_alias(style)
        size = max(5.0, float(style.get("fontSize") or 11))
        line_height = size * max(1.05, float(style.get("lineHeight") or 1.15))
        lines = _wrap_text(text, width - margin * 2, font, size)
        required = max(line_height, len(lines) * line_height) + 8
        ensure_space(required)
        align = {"center": 1, "right": 2, "justify": 3}.get(str(style.get("align") or "left"), 0)
        rect = pymupdf.Rect(margin, y, width - margin, y + required)
        page.insert_textbox(rect, "\n".join(lines), fontname=font, fontsize=size, color=_rgb(style.get("color")), align=align, lineheight=float(style.get("lineHeight") or 1.15))
        y += required + (4 if block_type == "list_item" else 8)

    output = document.tobytes(garbage=4, deflate=True, clean=True)
    document.close()
    return output


def find_office_executable() -> str | None:
    explicit = os.environ.get("LUMINA_OFFICE")
    candidates = [explicit] if explicit else []
    candidates.extend(filter(None, [shutil.which("libreoffice"), shutil.which("soffice")]))
    if os.name == "nt":
        candidates.extend([
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ])
    for candidate in candidates:
        if candidate and Path(candidate).exists():
            return str(candidate)
    return None


def convert_docx_bytes_to_pdf(docx_bytes: bytes) -> tuple[bytes | None, str | None]:
    executable = find_office_executable()
    if not executable:
        return None, None
    with tempfile.TemporaryDirectory(prefix="lumina-docx-") as directory:
        input_path = Path(directory) / "document.docx"
        input_path.write_bytes(docx_bytes)
        try:
            completed = subprocess.run(
                [executable, "--headless", "--convert-to", "pdf", "--outdir", directory, str(input_path)],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=90,
                check=False,
            )
        except (OSError, subprocess.SubprocessError):
            return None, None
        output_path = Path(directory) / "document.pdf"
        if completed.returncode == 0 and output_path.exists() and output_path.stat().st_size > 100:
            return output_path.read_bytes(), "libreoffice"
    return None, None


def render_document_model(model: dict[str, Any], prefer_office: bool = True) -> dict[str, Any]:
    docx_bytes = model_to_docx_bytes(model)
    pdf_bytes: bytes | None = None
    converter = "pymupdf-reflow"
    if prefer_office:
        pdf_bytes, office_converter = convert_docx_bytes_to_pdf(docx_bytes)
        if pdf_bytes:
            converter = office_converter or "office"
    if not pdf_bytes:
        pdf_bytes = model_to_fallback_pdf_bytes(model)
    return {
        "docx": docx_bytes,
        "pdf": pdf_bytes,
        "markdown": model_to_markdown(model),
        "converter": converter,
    }


__all__ = [
    "MODEL_VERSION",
    "pdf_to_document_model",
    "model_to_docx_bytes",
    "model_to_markdown",
    "model_to_fallback_pdf_bytes",
    "render_document_model",
    "find_office_executable",
]
